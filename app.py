from io import BytesIO
from pathlib import Path
import re

from flask import Flask, jsonify, render_template, request, send_file
from docx import Document
from security_logic import determine_level, get_base_requirements, get_measures_for_level

app = Flask(__name__)


def safe_filename(name: str) -> str:
    cleaned = "".join(ch for ch in (name or "") if ch.isalnum() or ch in (" ", "-", "_", "."))
    cleaned = cleaned.strip() or "Отчет ИСПДн"
    if not cleaned.lower().endswith(".docx"):
        cleaned += ".docx"
    return cleaned


def act_context_from_answers(answers: dict) -> dict:
    data_type = answers.get("dataType")
    threats_raw = answers.get("threats", [])
    employees_only = answers.get("employeesOnly", False)
    non_employee_scope = answers.get("nonEmployeeScope")

    if not employees_only and not non_employee_scope:
        non_employee_scope = "over_100k"

    threats_clean = [t for t in threats_raw if t and t != "unknown"]
    data_type_map = {
        "special": "специальные категории персональных данных",
        "biometric": "биометрические персональные данные",
        "public": "общедоступные персональные данные",
        "other": "иные категории персональных данных",
    }
    threats_map = {"1": "Угрозы 1-го типа", "2": "Угрозы 2-го типа", "3": "Угрозы 3-го типа"}

    categories_text = data_type_map.get(data_type, "Категория не указана")

    if threats_clean:
        threats_text = ", ".join(threats_map.get(t, t) for t in threats_clean)
    elif "unknown" in threats_raw:
        threats_text = "Тип актуальных угроз не определен (выбран «не известен»)"
    else:
        threats_text = "Тип актуальных угроз не указан"

    subjects_text = "Только сотрудники" if employees_only else "Не только сотрудники"

    if non_employee_scope == "over_100k":
        volume_text = "Более 100 000"
    elif non_employee_scope == "under_100k":
        volume_text = "До 100 000"
    else:
        volume_text = "—"

    level_value = determine_level(data_type, threats_clean, employees_only, non_employee_scope) if threats_clean else None
    level_map = {1: "1-го", 2: "2-го", 3: "3-го", 4: "4-го"}
    level_text = level_map.get(level_value, "—")

    return {
        "categories_text": categories_text,
        "threats_text": threats_text,
        "subjects_text": subjects_text,
        "volume_text": volume_text,
        "level_text": level_text,
    }


def replace_underscores(paragraph, value: str, remove_hints: bool = True) -> None:
    paragraph.text = re.sub(r"_+", value, paragraph.text)
    if remove_hints:
        paragraph.text = re.sub(r"\s*\([^)]*\)", "", paragraph.text)


def replace_placeholders(doc: Document, replacements: dict) -> None:
    def replace_text(text: str) -> str:
        new_text = text
        for key, val in replacements.items():
            placeholder = f"[{key}]"
            new_text = new_text.replace(placeholder, val)
        return new_text

    def process_paragraph(paragraph) -> None:
        if not paragraph.runs:
            return
        original = "".join(run.text for run in paragraph.runs)
        new_text = replace_text(original)
        if new_text != original:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""

    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)


def generate_act_document(
    template_path: Path,
    org_name: str,
    head_position: str,
    system_name: str,
    context: dict,
) -> BytesIO:
    doc = Document(template_path)

    replacements = {
        "Должность руководителя организации": head_position or "—",
        "Наименование организации": org_name or "—",
        "Наименование ИСПДн": system_name or "—",
        "Категория персональных данных": context["categories_text"],
        "Только сотрудники/Не только сотрудники": context["subjects_text"],
        "Объём обрабатываемых субъектов": context["volume_text"],
        "Тип актуальных угроз": context["threats_text"],
        "Вычисленный уровень защищённости ИСПДн": context["level_text"],
    }

    replace_placeholders(doc, replacements)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/evaluate", methods=["POST"])
def evaluate():
    payload = request.get_json(force=True)
    data_type = payload.get("dataType")
    threats = payload.get("threats", [])
    employees_only = payload.get("employeesOnly", False)
    non_employee_scope = payload.get("nonEmployeeScope")

    if not employees_only and not non_employee_scope:
        non_employee_scope = "over_100k"

    unknown_threats = "unknown" in threats

    def build_levels_for_unknown():
        possibilities = []
        for threat_type in ("1", "2", "3"):
            lvl = determine_level(data_type, [threat_type], employees_only, non_employee_scope)
            possibilities.append({"threatType": threat_type, "level": lvl})
        return possibilities

    if unknown_threats:
        possible_levels = build_levels_for_unknown()
        max_level = max((item["level"] for item in possible_levels), default=4)
        return jsonify(
            {
                "level": max_level,
                "possibleLevels": possible_levels,
                "unknownThreats": True,
                "baseRequirements": [],
                "measures": [],
            }
        )

    level = determine_level(data_type, threats, employees_only, non_employee_scope)
    base_requirements = get_base_requirements(level)
    measures = get_measures_for_level(level)

    return jsonify(
        {
            "level": level,
            "baseRequirements": base_requirements,
            "measures": [
                {
                    "code": measure.code,
                    "section": measure.section,
                    "description": measure.description,
                }
                for measure in measures
            ],
        }
    )


@app.route("/export", methods=["POST"])
def export_report():
    payload = request.get_json(force=True) or {}
    file_name = safe_filename(payload.get("fileName", "Отчет ИСПДн"))
    answers = payload.get("payload", {})

    data_type = answers.get("dataType")
    threats = answers.get("threats", [])
    employees_only = answers.get("employeesOnly", False)
    non_employee_scope = answers.get("nonEmployeeScope")

    if not employees_only and not non_employee_scope:
        non_employee_scope = "over_100k"

    unknown_threats = "unknown" in threats

    def build_levels_for_unknown():
        possibilities = []
        for threat_type in ("1", "2", "3"):
            lvl = determine_level(data_type, [threat_type], employees_only, non_employee_scope)
            possibilities.append({"threatType": threat_type, "level": lvl})
        return possibilities

    doc = Document()
    doc.add_heading("Отчет по защищенности ИСПДн", level=1)

    if unknown_threats:
        possible_levels = build_levels_for_unknown()
        doc.add_paragraph("Тип актуальных угроз выбран как «не известен».")
        doc.add_heading("Возможные уровни защищенности", level=2)
        for item in possible_levels:
            doc.add_paragraph(
                f"Угрозы {item['threatType']} типа: уровень {item['level']}", style="List Bullet"
            )
        doc.add_paragraph(
            "Для уточнения уровня защищенности закажите услугу специалиста по определению типа актуальных угроз, "
            "после этого выполните перерасчет."
        )
    else:
        level = determine_level(data_type, threats, employees_only, non_employee_scope)
        base_requirements = get_base_requirements(level)
        measures = get_measures_for_level(level)

        doc.add_paragraph(f"Уровень защищенности: {level}")
        doc.add_paragraph("")

        doc.add_heading("Организационные требования", level=2)
        for req in base_requirements:
            doc.add_paragraph(req, style="List Bullet")

        doc.add_heading("Базовый набор мер", level=2)
        for measure in measures:
            doc.add_paragraph(f"{measure.code} — {measure.section}", style="List Bullet")
            doc.add_paragraph(measure.description)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=file_name,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.route("/export-act", methods=["POST"])
def export_act():
    payload = request.get_json(force=True) or {}
    file_name = safe_filename(payload.get("fileName", "АКТ_уровень_защищенности.docx"))
    answers = payload.get("payload", {})
    user_inputs = payload.get("userInputs", {})

    org_name = (user_inputs.get("organization") or "").strip()
    head_position = (user_inputs.get("headPosition") or "").strip()
    system_name = (user_inputs.get("systemName") or "").strip()
    if not org_name or not head_position or not system_name:
        return jsonify({"error": "Не заполнены обязательные поля для акта"}), 400

    template_path = Path("act.docx")
    if not template_path.exists():
        return jsonify({"error": "Шаблон акта не найден"}), 404

    threats_raw = answers.get("threats", [])
    if "unknown" in threats_raw:
        return jsonify({"error": "Тип угроз не определен, используйте ручной шаблон"}), 400

    if not answers.get("nonEmployeeScope"):
        return jsonify({"error": "Не указан объём обрабатываемых субъектов"}), 400

    context = act_context_from_answers(answers)
    buffer = generate_act_document(template_path, org_name, head_position, system_name, context)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=file_name,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.route("/act-template", methods=["GET"])
def act_template():
    template_path = Path("act.docx")
    if not template_path.exists():
        return jsonify({"error": "Шаблон акта не найден"}), 404
    return send_file(
        template_path,
        as_attachment=True,
        download_name="АКТ_шаблон.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.route("/act-self", methods=["GET"])
def act_self():
    template_path = Path("act_self_work.docx")
    if not template_path.exists():
        return jsonify({"error": "Шаблон акта для ручного заполнения не найден"}), 404
    return send_file(
        template_path,
        as_attachment=True,
        download_name="АКТ_для_ручного_заполнения.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)
